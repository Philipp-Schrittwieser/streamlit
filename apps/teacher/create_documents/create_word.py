import io
from docx import Document

def add_formatted_text_open_creator(doc, text):
    lines = text.split('\n')
    current_paragraph = None
    
    for line in lines:
        line = line.strip()  # Entfernt Leerzeichen am Anfang und Ende
        if not line:  # Ignoriert leere Zeilen
            continue
        
        if line.startswith('# '):
            heading = doc.add_heading(level=1)
            add_formatted_run(heading, line[2:])
            current_paragraph = None
        elif line.startswith('## '):
            heading = doc.add_heading(level=2)
            add_formatted_run(heading, line[3:])
            current_paragraph = None
        elif line.startswith('### '):
            heading = doc.add_heading(level=3)
            add_formatted_run(heading, line[4:])
            current_paragraph = None
        elif line.startswith('#### '):
            heading = doc.add_heading(level=4)
            add_formatted_run(heading, line[5:])
            current_paragraph = None
        elif line.startswith('- ') or line.startswith('•\t'):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_run(p, line[2:])
            current_paragraph = None
        else:
            if current_paragraph is None:
                current_paragraph = doc.add_paragraph()
            else:
                current_paragraph.add_run('\n')  # Fügt einen Zeilenumbruch innerhalb des Absatzes hinzu
            add_formatted_run(current_paragraph, line)

def add_formatted_text_genius_ai(doc, text):
    lines = text.split('\n')
    current_paragraph = None
    
    for line in lines:
        line = line.strip()  # Entfernt Leerzeichen am Anfang und Ende
        if not line:  # Ignoriert leere Zeilen
            continue
        
        if line.startswith('# '):
            heading = doc.add_heading(level=1)
            add_formatted_run(heading, line[2:])
            current_paragraph = None
        elif line.startswith('## '):
            heading = doc.add_heading(level=1)
            add_formatted_run(heading, line[3:])
            current_paragraph = None
        elif line.startswith('### '):
            heading = doc.add_heading(level=2)
            add_formatted_run(heading, line[4:])
            current_paragraph = None
        elif line.startswith('#### '):
            heading = doc.add_heading(level=3)
            add_formatted_run(heading, line[5:])
            current_paragraph = None
        elif line.startswith('##### '):
            heading = doc.add_heading(level=4)
            add_formatted_run(heading, line[6:])
            current_paragraph = None
        elif line.startswith('* ') or line.startswith('•\t'):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_run(p, line[2:])
            current_paragraph = None
        else:
            if current_paragraph is None:
                current_paragraph = doc.add_paragraph()
            else:
                current_paragraph.add_run('\n\n')  # Fügt einen Zeilenumbruch innerhalb des Absatzes hinzu
            add_formatted_run(current_paragraph, line)

def add_formatted_run(paragraph, text):
    parts = text.split('**')
    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        if i % 2 == 1:  # Ungerade Indizes sind fett gedruckt
            run.bold = True

def create_document_for_genius_ai(content):
  doc = Document()
  add_formatted_text_genius_ai(doc, content)
  bio = io.BytesIO()
  doc.save(bio)
  return bio


def create_combined__exercise_document(lesetext, fragen, lösungen):
    doc = Document()
    add_formatted_text_genius_ai(doc, lesetext)
    doc.add_page_break()
    add_formatted_text_genius_ai(doc, fragen)
    doc.add_page_break()
    add_formatted_text_genius_ai(doc, lösungen)
    bio = io.BytesIO()
    doc.save(bio)
    return bio

def create_combined_grammar_document(exercises, solutions):
    doc = Document()
    add_formatted_text_genius_ai(doc, exercises)
    doc.add_page_break()
    add_formatted_text_genius_ai(doc, solutions)
    bio = io.BytesIO()
    doc.save(bio)
    return bio


def format_questions_and_answers(topic, qa_pairs, lines):
    """
    Formatiert Fragen und Antworten für Word/Streamlit
    
    Args:
        topic: String mit Thema
        qa_pairs: Liste mit Frage-Antwort Paaren in verschiedenen Formaten:
            - [{"question": "...", "answer": "..."}, ...]
            - [{"exercise": "...", "solution": "..."}, ...]
    """

    fragen = f"## {topic} - Aufgaben:\n\n"
    lösungen = f"## {topic} - Lösungen:\n\n"

    for i, qa in enumerate(qa_pairs, 1):
        question = qa.get("question") or qa.get("exercise") 
        answer = qa.get("answer") or qa.get("solution")

        # Bei Gaptext brauch ich keine Antwortzeile ("____")
        # Das u200B ist in Word ein leeres Zeichen (sieht man nicht)
        # "\n\n" springt in die nächste Zeile
        # Wenn 3 "___" hintereinander vorkommen, dann ist es wsl von den Blanks aus den Übungen
        
        if "___" in question:
            question = question.replace("___", "_________")
            fragen += f"### **{i}. {question}**\n\n\u200B\n\n"
            
        # Wenn nicht halte dich an die Lines
        else: fragen += f"### **{i}. {question}**{lines}"
        lösungen += f"### **{i}. {question}**\n\n{answer}\n\n"

    return fragen, lösungen
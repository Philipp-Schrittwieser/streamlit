import streamlit as st
import time
import pandas as pd
from apps.teacher.llms.gpt.generate_1_text_gpt import generate_1_text_gpt
from apps.teacher.llms.gpt.generate_2_qas_gpt import generate_2_qas_gpt
import io
from docx import Document

# st.subheader("1️⃣. Lesetext mit KI generieren")
# st.subheader("2️⃣. Fragen mit KI generieren (Schüler)")
# st.subheader("3️⃣. Antworten mit KI generieren (Lehrer)")

def show_generate_popup(type):
  st.success(f"Generiere {type}... ⛏️")
  st.toast(f"KI generiert {type}... 💡")
  time.sleep(1)
  st.toast("Bitte warte einen kurzen Moment ⏰ ")

def show_finish_popup():
  st.balloons()
  time.sleep(0.75)
  st.rerun()

def add_formatted_text(doc, text):
    lines = text.split('\n')
    current_paragraph = None
    
    for line in lines:
        line = line.strip()  # Entfernt Leerzeichen am Anfang und Ende
        if not line:  # Ignoriert leere Zeilen
            continue
        
        if line.startswith('# '):
            heading = doc.add_heading(level=1)
            add_formatted_run(heading, line[2:])
            current_paragraph = None
        elif line.startswith('## '):
            heading = doc.add_heading(level=2)
            add_formatted_run(heading, line[3:])
            current_paragraph = None
        elif line.startswith('### '):
            heading = doc.add_heading(level=3)
            add_formatted_run(heading, line[4:])
            current_paragraph = None
        elif line.startswith('#### '):
            heading = doc.add_heading(level=4)
            add_formatted_run(heading, line[5:])
            current_paragraph = None
        elif line.startswith('- ') or line.startswith('•\t'):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_run(p, line[2:])
            current_paragraph = None
        else:
            if current_paragraph is None:
                current_paragraph = doc.add_paragraph()
            else:
                current_paragraph.add_run('\n')  # Fügt einen Zeilenumbruch innerhalb des Absatzes hinzu
            add_formatted_run(current_paragraph, line)

def add_formatted_run(paragraph, text):
    parts = text.split('**')
    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        if i % 2 == 1:  # Ungerade Indizes sind fett gedruckt
            run.bold = True

if "exercise_sheet_level" not in st.session_state:
  st.session_state.exercise_sheet_level = "1_text"

if "response" not in st.session_state:
  st.session_state.response = ""

left, right = st.columns([15, 1], gap="small", vertical_alignment="center")
left.title("Arbeitsblatt Generator 📝", anchor=False)

if st.session_state.exercise_sheet_level != "1_text":
  if right.button(":material/restart_alt:", use_container_width=False):
    st.session_state.exercise_sheet_level = "1_text"
    st.session_state.topic = ""
    st.session_state.response = ""
    st.session_state.qas = []
    st.session_state.fragen = ""
    st.session_state.lösungen = ""
    st.rerun()

st.text("")  # Fügt eine Leerzeile hinzu


if st.session_state.exercise_sheet_level == "1_text":
  one, two, three = st.columns(3)
  one.markdown("<h4>1️⃣. <u>Lesetext</u> 📖", unsafe_allow_html=True)
  two.markdown("<h4>2️⃣. Aufgaben", unsafe_allow_html=True)
  three.markdown("<h4>3️⃣. Download", unsafe_allow_html=True)

if st.session_state.exercise_sheet_level == "2_qas":
  one, two, three = st.columns(3)
  one.markdown("<h4>1️⃣. Lesetext", unsafe_allow_html=True)
  two.markdown("<h4>2️⃣. <u>Aufgaben</u> ⁉️", unsafe_allow_html=True)
  three.markdown("<h4>3️⃣. Download", unsafe_allow_html=True)

if st.session_state.exercise_sheet_level == "3_answers":
  one, two, three = st.columns(3)
  one.markdown("<h4>1️⃣. Lesetext", unsafe_allow_html=True)
  two.markdown("<h4>2️⃣. Aufgaben", unsafe_allow_html=True)
  three.markdown("<h4>3️⃣. <u>Download</u> 💾 ", unsafe_allow_html=True)

st.text("")


if st.session_state.exercise_sheet_level == "1_text":
  st.subheader("1. Lesetext einfügen oder generieren...", divider="red", anchor=False)
  st.write("Möchtest du einen Text einfügen, aus dem wir dein Arbeitsblatt erstellen (Möglichkeit A) oder sollen wir einen neuen Text für dich generieren (Möglichkeit B)?")
  st.write("Tippen dein Thema unterhalb ein und dann wähle Option A oder B.")

  st.text("")

  topic = st.text_input("Thema eintippen:", placeholder="z.B. Klimawandel")
  st.session_state.topic = topic

  st.text("")


  left, right = st.columns(2, gap="large")

  with left:
    st.subheader("Möglichkeit A", divider="violet")
    # st.write("##### Einfügen...")
    topic_text = st.text_area("Text hier einfügen:", placeholder="Hier steht dein Text...", height=68)
    st.write("")
    st.session_state.response = topic_text

    if st.button(label="Fortfahren :material/arrow_right_alt:", key="button-blue", use_container_width=True):
      
      if st.session_state.topic == "":
        st.toast("Thema fehlt :exclamation:")
        st.error("Bitte Thema eingeben :exclamation:")
        time.sleep(0.5)
      
      if st.session_state.response == "":
        st.toast("Text fehlt :exclamation:")
        st.error("Bitte Text einfügen :exclamation:")

      if st.session_state.topic != "" and st.session_state.response != "":
        st.session_state.response = topic_text
        st.session_state.exercise_sheet_level = "2_qas"
        show_finish_popup()


    with right:
      st.subheader("Möglichkeit B", divider="green")
      st.write("Text generieren lassen:")
      # st.write("##### Generieren...")

      if st.button(label="Lesetext generieren :material/laps:", key="button-blue2", use_container_width=True):
        
        if st.session_state.topic == "":
          st.toast("Thema fehlt :exclamation:")
          st.error("Bitte Thema eingeben :exclamation:")

        else:
          show_generate_popup("Lesetext")

          with st.spinner(''):
            ai_model = "Max Creator"
            if ai_model == "Max Creator":
              response = generate_1_text_gpt(topic)
            elif ai_model == "Genius AI":
              pass
              # response = generate_quiz_gemini(user_text, num_questions, time_limit)
            else:
              st.error("Kein gültiges KI-Modell ausgewählt :exclamation:")
              
            st.session_state.response = response
            st.session_state.exercise_sheet_level = "2_qas"
            # st.success("Fertig generiert 🎉🎉🎉")      
            show_finish_popup()



elif st.session_state.exercise_sheet_level == "2_qas":
  st.subheader("2. Aufgaben und Lösungen erstellen...", divider="violet", anchor=False)

  st.write("Dein Lesetext ist im Dropdown-Menü verfügbar. Klicke jetzt auf den Button darunter, um Aufgaben und Lösungen zu deinem Text zu erstellen.")

  with st.expander("Lesetext hier ansehen"):
    st.write(st.session_state.response)
  
  if st.button(label="Aufgaben generieren :material/laps:", key="button-blue"):
    show_generate_popup("Aufgaben und Lösungen")

    with st.spinner(''):
      ai_model = "Max Creator"
      if ai_model == "Max Creator":
        qas = generate_2_qas_gpt(st.session_state.response)
      elif ai_model == "Genius AI":
        pass
        # response = generate_quiz_gemini(user_text, num_questions, time_limit)
      else:
        st.error("Kein gültiges KI-Modell ausgewählt :exclamation:")
      
      st.session_state.qas = qas
      st.session_state.exercise_sheet_level = "3_answers"
      # st.success("Fertig generiert 🎉🎉🎉")      
      show_finish_popup()

elif st.session_state.exercise_sheet_level == "3_answers":
  st.subheader("3. Alle Unterlagen herunterladen...", divider="green", anchor=False)

  st.write("Hier kannst du noch einmal alle Dokumente einsehen und wenn du zufrieden bist einfach herunterladen.")

  with st.expander("Generierter Lesetext"):
    st.write(st.session_state.response)
  
  with st.expander("Generierte Aufgaben"):
    fragen = f"## {st.session_state.topic} - Aufgaben:\n\n"
    lösungen = f"## {st.session_state.topic} - Lösungen:\n\n" # inkludiert Fragen und Lösungen

    for i, qa in enumerate(st.session_state.qas, 1):  # Start bei 1 statt 0
      st.markdown(f"### **{i}. {qa['question']}**")
      st.markdown(f"{qa['answer']}")

      fragen += f"### **{i}. {qa['question']}**\n\n\u200B\n\n_________________________________________________________________________________________________________\n\n\u200B\n\n_________________________________________________________________________________________________________\n\n\u200B\n\n"
      st.session_state.fragen = fragen
      lösungen += f"### **{i}. {qa['question']}**\n\n{qa['answer']}\n\n"
      st.session_state.lösungen = lösungen

  # Funktion zum Erstellen eines Dokuments und Speichern in BytesIO
  def create_document(content, filename):
      doc = Document()
      add_formatted_text(doc, content)
      bio = io.BytesIO()
      doc.save(bio)
      return bio, filename

  # Erstellen Sie die verschiedenen Dokumente
  lesetext_bio, lesetext_filename = create_document(st.session_state.response, f"{st.session_state.topic}_Lesetext.docx")
  fragen_bio, fragen_filename = create_document(st.session_state.fragen, f"{st.session_state.topic}_Fragen.docx")
  lösungen_bio, lösungen_filename = create_document(st.session_state.lösungen, f"{st.session_state.topic}_Lösungen.docx")

  # Erstellen Sie ein kombiniertes Dokument
  combined_doc = Document()
  add_formatted_text(combined_doc, st.session_state.response)
  combined_doc.add_page_break()  # Neue Seite vor Fragen
  add_formatted_text(combined_doc, st.session_state.fragen)
  combined_doc.add_page_break()  # Neue Seite vor Lösungen
  add_formatted_text(combined_doc, st.session_state.lösungen)
  combined_bio = io.BytesIO()
  combined_doc.save(combined_bio)

  with st.expander("Einzelne Dokumente"):
    # Erstellen Sie die Download-Buttons
    one, two, three = st.columns(3)

    one.download_button(
        use_container_width=True,
        label="Lesetext :material/download:",
        data=lesetext_bio.getvalue(),
        file_name=lesetext_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    two.download_button(
        use_container_width=True,
        label="Fragen :material/download:",
        data=fragen_bio.getvalue(),
        file_name=fragen_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    three.download_button(
        use_container_width=True,
        label="Lösungen :material/download:",
        data=lösungen_bio.getvalue(),
        file_name=lösungen_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

  st.download_button(
      label="Gesamtes Dokument herunterladen :material/download:",
      data=combined_bio.getvalue(),
      file_name=f"{st.session_state.topic}_Alle_Unterlagen.docx",
      mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
  )